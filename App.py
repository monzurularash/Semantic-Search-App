import os
import numpy as np
import hashlib
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QLabel, QFileDialog, QTextEdit, QProgressBar

# Load the model
model_name = 'all-MiniLM-L6-v2'  # Replace with any pre-trained model you prefer
model = SentenceTransformer(model_name)

# Function to calculate the hash of a directory's contents
def calculate_directory_hash(directory_path):
    hash_md5 = hashlib.md5()
    for root, dirs, files in os.walk(directory_path):
        for file in sorted(files):
            if file != 'dir_hash.txt':  # Skip the hash file itself
                file_path = os.path.join(root, file)
                with open(file_path, "rb") as f:
                    for chunk in iter(lambda: f.read(4096), b""):
                        hash_md5.update(chunk)
    return hash_md5.hexdigest()

# Function to load and encode .docx and .txt files
def load_and_encode(directory_path, progress_callback=None):
    encoded_data = []
    file_names = []
    
    file_list = []
    for root, dirs, files in os.walk(directory_path):
        for filename in files:
            if filename.endswith(".docx") or filename.endswith(".txt"):
                file_list.append(os.path.join(root, filename))
    
    total_files = len(file_list)
    for i, file_path in enumerate(file_list):
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding='utf-8') as f:
                full_text = f.read()
        
        encoded_text = model.encode(full_text, convert_to_tensor=False)
        encoded_data.append(encoded_text)
        file_names.append(file_path)
        
        if progress_callback:
            progress_callback(i + 1, total_files)
    
    return np.array(encoded_data), np.array(file_names)

# Function to load embeddings and file names if the directory has not changed
def load_data_if_unchanged(directory_path, progress_callback=None):
    current_hash = calculate_directory_hash(directory_path)
    hash_file_path = os.path.join(os.path.dirname(directory_path), 'dir_hash.txt')

    if os.path.exists(hash_file_path):
        with open(hash_file_path, 'r') as hash_file:
            saved_hash = hash_file.read()
        
        if saved_hash == current_hash:
            print("Directory unchanged. Loading existing data.")
            return np.load(os.path.join(directory_path, 'text_embeddings.npy')), \
                   np.load(os.path.join(directory_path, 'file_names.npy'))
    
    print("Directory changed or first time loading. Encoding data.")
    text_embeddings, file_names = load_and_encode(directory_path, progress_callback)
    np.save(os.path.join(directory_path, 'text_embeddings.npy'), text_embeddings)
    np.save(os.path.join(directory_path, 'file_names.npy'), file_names)

    with open(hash_file_path, 'w') as hash_file:
        hash_file.write(current_hash)
    
    return text_embeddings, file_names

# Function to query the model
def query_model(query, text_embeddings):
    query_embedding = model.encode([query], convert_to_tensor=False)
    similarities = cosine_similarity(query_embedding, text_embeddings)[0]
    return similarities

# Function to run the UI
def run_app():
    app = QApplication([])

    window = QWidget()
    window.setWindowTitle("Document Search")

    layout = QVBoxLayout()

    directory_label = QLabel("Select Directory:")
    layout.addWidget(directory_label)

    directory_button = QPushButton("Choose Directory")
    layout.addWidget(directory_button)

    query_label = QLabel("Enter your query:")
    layout.addWidget(query_label)

    query_text = QTextEdit()
    layout.addWidget(query_text)

    search_button = QPushButton("Search")
    layout.addWidget(search_button)

    progress_bar = QProgressBar()
    layout.addWidget(progress_bar)

    result_label = QLabel("Results:")
    layout.addWidget(result_label)

    def choose_directory():
        directory = QFileDialog.getExistingDirectory(window, "Select Directory")
        if directory:
            directory_label.setText(f"Selected Directory: {directory}")
            window.directory_path = directory

    def update_progress(current, total):
        progress_bar.setValue(int((current / total) * 100))

    def search_documents():
        query = query_text.toPlainText()
        directory_path = window.directory_path

        text_embeddings, file_names = load_data_if_unchanged(directory_path, update_progress)
        similarities = query_model(query, text_embeddings)

        top_indices = similarities.argsort()[-5:][::-1]
        results = ""
        for index in top_indices:
            doc_path = file_names[index]
            if doc_path.endswith(".docx"):
                doc = Document(doc_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
            elif doc_path.endswith(".txt"):
                with open(doc_path, "r", encoding='utf-8') as f:
                    full_text = f.read()
            preview = full_text[:500]  # Preview of the first 500 characters
            results += f"File: {doc_path}\nPreview: {preview}\n\n"

        result_label.setText(results)

    directory_button.clicked.connect(choose_directory)
    search_button.clicked.connect(search_documents)

    window.setLayout(layout)
    window.show()
    app.exec_()

if __name__ == "__main__":
    run_app()
